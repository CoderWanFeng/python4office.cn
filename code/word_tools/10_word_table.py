#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word表格处理工具
功能：批量处理Word文档中的表格，支持数据提取、样式修改和表格操作
"""

import sys
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from PySide6.QtWidgets import (QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, 
                              QPushButton, QTableWidget, QTableWidgetItem, 
                              QHeaderView, QTabWidget, QComboBox, QRadioButton,
                              QButtonGroup, QSpinBox, QCheckBox, QMessageBox)
from PySide6.QtCore import QTimer

# 添加父目录到路径，以便导入common模块
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
sys.path.append(parent_dir)

from common.ui_base import BaseMainWindow, FileSelectionWidget, OperationWidget


class WordTableWindow(BaseMainWindow):
    """Word表格处理工具主窗口"""
    
    def __init__(self):
        super().__init__("Word表格处理工具", 1000, 700)
        
    def setup_content(self):
        """设置内容区域"""
        # 文件选择组件
        self.file_selector = FileSelectionWidget("选择要处理表格的Word文件", multi_selection=True)
        self.content_layout.addWidget(self.file_selector)
        
        # 使用选项卡组织界面
        self.tab_widget = QTabWidget()
        self.content_layout.addWidget(self.tab_widget)
        
        # 表格提取选项卡
        self.extract_tab = QWidget()
        self.setup_extract_tab()
        self.tab_widget.addTab(self.extract_tab, "表格提取")
        
        # 表格样式选项卡
        self.style_tab = QWidget()
        self.setup_style_tab()
        self.tab_widget.addTab(self.style_tab, "表格样式")
        
        # 表格操作选项卡
        self.operation_tab = QWidget()
        self.setup_operation_tab()
        self.tab_widget.addTab(self.operation_tab, "表格操作")
        
        # 操作区域
        self.operation = OperationWidget("表格处理操作")
        self.operation.start_operation = self.process_tables
        self.operation.signals.progress.connect(self.update_progress)
        self.operation.signals.finished.connect(self.table_processing_finished)
        self.content_layout.addWidget(self.operation)
        
    def setup_extract_tab(self):
        """设置表格提取选项卡"""
        layout = QVBoxLayout(self.extract_tab)
        
        # 提取选项
        extract_group = QWidget()
        extract_layout = QVBoxLayout(extract_group)
        extract_layout.addWidget(QLabel("提取选项:"))
        
        # 输出格式
        format_layout = QHBoxLayout()
        format_layout.addWidget(QLabel("输出格式:"))
        self.output_format_combo = QComboBox()
        self.output_format_combo.addItems(["Excel文件", "CSV文件", "JSON文件"])
        format_layout.addWidget(self.output_format_combo)
        extract_layout.addLayout(format_layout)
        
        # 合并方式
        merge_layout = QHBoxLayout()
        merge_layout.addWidget(QLabel("表格合并方式:"))
        self.merge_combo = QComboBox()
        self.merge_combo.addItems(["所有表格合并", "每个表格单独保存", "按文档分组保存"])
        merge_layout.addWidget(self.merge_combo)
        extract_layout.addLayout(merge_layout)
        
        layout.addWidget(extract_group)
        
        # 预览区域
        layout.addWidget(QLabel("表格预览:"))
        self.preview_table = QTableWidget()
        self.preview_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.preview_table.setMaximumHeight(200)
        layout.addWidget(self.preview_table)
        
        # 预览按钮
        self.preview_btn = QPushButton("预览第一个文档中的表格")
        self.preview_btn.clicked.connect(self.preview_tables)
        layout.addWidget(self.preview_btn)
        
    def setup_style_tab(self):
        """设置表格样式选项卡"""
        layout = QVBoxLayout(self.style_tab)
        
        # 样式选项
        style_group = QWidget()
        style_layout = QVBoxLayout(style_group)
        style_layout.addWidget(QLabel("样式修改选项:"))
        
        # 表格对齐
        align_layout = QHBoxLayout()
        align_layout.addWidget(QLabel("表格对齐:"))
        self.table_align_combo = QComboBox()
        self.table_align_combo.addItems(["左对齐", "居中", "右对齐"])
        align_layout.addWidget(self.table_align_combo)
        style_layout.addLayout(align_layout)
        
        # 表格宽度
        width_layout = QHBoxLayout()
        width_layout.addWidget(QLabel("表格宽度(%):"))
        self.table_width_spin = QSpinBox()
        self.table_width_spin.setRange(10, 100)
        self.table_width_spin.setValue(100)
        width_layout.addWidget(self.table_width_spin)
        style_layout.addLayout(width_layout)
        
        # 边框设置
        border_layout = QHBoxLayout()
        border_layout.addWidget(QLabel("边框宽度(磅):"))
        self.border_width_spin = QSpinBox()
        self.border_width_spin.setRange(0, 12)
        self.border_width_spin.setValue(1)
        border_layout.addWidget(self.border_width_spin)
        style_layout.addLayout(border_layout)
        
        # 表头样式
        header_layout = QVBoxLayout()
        header_layout.addWidget(QLabel("表头样式:"))
        
        self.header_bold_checkbox = QCheckBox("表头加粗")
        self.header_bold_checkbox.setChecked(True)
        header_layout.addWidget(self.header_bold_checkbox)
        
        header_align_layout = QHBoxLayout()
        header_align_layout.addWidget(QLabel("表头对齐:"))
        self.header_align_combo = QComboBox()
        self.header_align_combo.addItems(["左对齐", "居中", "右对齐"])
        self.header_align_combo.setCurrentIndex(1)  # 默认居中
        header_align_layout.addWidget(self.header_align_combo)
        header_layout.addLayout(header_align_layout)
        
        style_layout.addLayout(header_layout)
        
        layout.addWidget(style_group)
        
    def setup_operation_tab(self):
        """设置表格操作选项卡"""
        layout = QVBoxLayout(self.operation_tab)
        
        # 操作选项
        operation_group = QWidget()
        operation_layout = QVBoxLayout(operation_group)
        operation_layout.addWidget(QLabel("表格操作选项:"))
        
        # 操作类型
        self.operation_group = QButtonGroup()
        
        self.add_col_radio = QRadioButton("添加列")
        self.operation_group.addButton(self.add_col_radio, 0)
        operation_layout.addWidget(self.add_col_radio)
        
        self.add_row_radio = QRadioButton("添加行")
        self.operation_group.addButton(self.add_row_radio, 1)
        operation_layout.addWidget(self.add_row_radio)
        
        self.del_col_radio = QRadioButton("删除列")
        self.operation_group.addButton(self.del_col_radio, 2)
        operation_layout.addWidget(self.del_col_radio)
        
        self.del_row_radio = QRadioButton("删除行")
        self.operation_group.addButton(self.del_row_radio, 3)
        operation_layout.addWidget(self.del_row_radio)
        
        layout.addWidget(operation_group)
        
        # 操作参数
        param_group = QWidget()
        param_layout = QVBoxLayout(param_group)
        param_layout.addWidget(QLabel("操作参数:"))
        
        # 位置设置
        position_layout = QHBoxLayout()
        position_layout.addWidget(QLabel("位置:"))
        self.position_spin = QSpinBox()
        self.position_spin.setMinimum(1)
        self.position_spin.setValue(1)
        position_layout.addWidget(self.position_spin)
        param_layout.addLayout(position_layout)
        
        # 数量设置
        count_layout = QHBoxLayout()
        count_layout.addWidget(QLabel("数量:"))
        self.count_spin = QSpinBox()
        self.count_spin.setMinimum(1)
        self.count_spin.setValue(1)
        count_layout.addWidget(self.count_spin)
        param_layout.addLayout(count_layout)
        
        layout.addWidget(param_group)
        
        # 连接信号
        self.operation_group.buttonToggled.connect(self.toggle_operation_params)
        
    def preview_tables(self):
        """预览表格"""
        files = self.file_selector.get_files()
        if not files:
            self.show_warning("请先选择Word文件")
            return
            
        try:
            # 加载第一个文档
            doc = Document(files[0])
            
            # 查找第一个表格
            if doc.tables:
                table = doc.tables[0]
                
                # 设置预览表格
                self.preview_table.setRowCount(len(table.rows))
                self.preview_table.setColumnCount(len(table.columns))
                
                # 填充数据
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        text = cell.text.strip()
                        self.preview_table.setItem(i, j, QTableWidgetItem(text))
                
                self.statusBar().showMessage(f"预览: {os.path.basename(files[0])} 第1个表格")
            else:
                self.show_warning("文档中没有找到表格")
                
        except Exception as e:
            self.show_error(f"预览表格失败: {str(e)}")
            
    def toggle_operation_params(self, button, checked):
        """根据操作类型切换参数"""
        if not checked:
            return
            
        operation_id = self.operation_group.checkedId()
        # 这里可以根据不同的操作类型启用/禁用某些参数
        if operation_id in [0, 1]:  # 添加列或行
            self.position_spin.setEnabled(True)
            self.count_spin.setEnabled(True)
        elif operation_id in [2, 3]:  # 删除列或行
            self.position_spin.setEnabled(True)
            self.count_spin.setEnabled(True)
        else:
            self.position_spin.setEnabled(False)
            self.count_spin.setEnabled(False)
            
    def process_tables(self):
        """处理表格"""
        files = self.file_selector.get_files()
        if not files:
            self.show_warning("请先选择Word文件")
            return
            
        # 获取当前选项卡
        current_tab = self.tab_widget.currentIndex()
        
        if current_tab == 0:  # 表格提取
            self.extract_tables(files)
        elif current_tab == 1:  # 表格样式
            self.style_tables(files)
        elif current_tab == 2:  # 表格操作
            self.operate_tables(files)
            
    def extract_tables(self, files):
        """提取表格"""
        try:
            # 获取提取选项
            output_format = self.output_format_combo.currentText()
            merge_mode = self.merge_combo.currentText()
            
            # 获取保存路径
            if output_format == "Excel文件":
                default_ext = ".xlsx"
                filter_str = "Excel文件 (*.xlsx)"
            elif output_format == "CSV文件":
                default_ext = ".csv"
                filter_str = "CSV文件 (*.csv)"
            else:  # JSON
                default_ext = ".json"
                filter_str = "JSON文件 (*.json)"
                
            save_dir = self.get_directory("选择表格保存目录")
            if not save_dir:
                return
                
            # 开始处理
            self.operation.set_running_state(True)
            self.operation.signals.progress.emit(0, 100, "开始提取表格...")
            
            # 使用定时器异步执行
            QTimer.singleShot(100, lambda: self.process_table_extraction(
                files, output_format, merge_mode, save_dir, default_ext
            ))
            
        except Exception as e:
            self.show_error(f"提取表格出错: {str(e)}")
            
    def style_tables(self, files):
        """修改表格样式"""
        try:
            # 获取样式选项
            table_align = self.table_align_combo.currentText()
            table_width = self.table_width_spin.value()
            border_width = self.border_width_spin.value()
            header_bold = self.header_bold_checkbox.isChecked()
            header_align = self.header_align_combo.currentText()
            
            # 获取保存目录
            save_dir = self.get_directory("选择处理后的文件保存目录")
            if not save_dir:
                return
                
            # 开始处理
            self.operation.set_running_state(True)
            self.operation.signals.progress.emit(0, 100, "开始修改表格样式...")
            
            # 使用定时器异步执行
            QTimer.singleShot(100, lambda: self.process_table_styling(
                files, table_align, table_width, border_width, header_bold, header_align, save_dir
            ))
            
        except Exception as e:
            self.show_error(f"修改表格样式出错: {str(e)}")
            
    def operate_tables(self, files):
        """操作表格"""
        try:
            # 获取操作选项
            operation_id = self.operation_group.checkedId()
            if operation_id == -1:
                self.show_warning("请选择一种表格操作")
                return
                
            position = self.position_spin.value()
            count = self.count_spin.value()
            
            # 获取保存目录
            save_dir = self.get_directory("选择处理后的文件保存目录")
            if not save_dir:
                return
                
            # 开始处理
            self.operation.set_running_state(True)
            self.operation.signals.progress.emit(0, 100, "开始表格操作...")
            
            # 使用定时器异步执行
            QTimer.singleShot(100, lambda: self.process_table_operations(
                files, operation_id, position, count, save_dir
            ))
            
        except Exception as e:
            self.show_error(f"表格操作出错: {str(e)}")
            
    def process_table_extraction(self, files, output_format, merge_mode, save_dir, default_ext):
        """处理表格提取"""
        try:
            total_files = len(files)
            all_data = []
            
            for i, file_path in enumerate(files):
                if not self.operation.is_running:
                    break
                    
                self.operation.signals.progress.emit(
                    int(i / total_files * 80), 
                    100, 
                    f"正在提取表格: {os.path.basename(file_path)}"
                )
                
                doc = Document(file_path)
                file_data = []
                
                for table_idx, table in enumerate(doc.tables):
                    # 提取表格数据
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    
                    file_data.append({
                        "filename": os.path.basename(file_path),
                        "table_index": table_idx + 1,
                        "data": table_data
                    })
                
                all_data.extend(file_data)
            
            # 保存数据
            self.operation.signals.progress.emit(80, 100, "正在保存提取的数据...")
            
            if output_format == "Excel文件":
                self.save_to_excel(all_data, merge_mode, save_dir, default_ext)
            elif output_format == "CSV文件":
                self.save_to_csv(all_data, merge_mode, save_dir, default_ext)
            else:  # JSON
                self.save_to_json(all_data, merge_mode, save_dir, default_ext)
            
            # 完成
            if self.operation.is_running:
                message = f"表格提取完成，共处理 {len(all_data)} 个表格"
                self.operation.signals.finished.emit(True, message)
            else:
                self.operation.signals.finished.emit(False, "用户取消操作")
                
        except Exception as e:
            self.operation.signals.finished.emit(False, f"提取表格过程中出错: {str(e)}")
            
    def process_table_styling(self, files, table_align, table_width, border_width, header_bold, header_align, save_dir):
        """处理表格样式"""
        try:
            total_files = len(files)
            
            for i, file_path in enumerate(files):
                if not self.operation.is_running:
                    break
                    
                self.operation.signals.progress.emit(
                    int(i / total_files * 90), 
                    100, 
                    f"正在处理表格: {os.path.basename(file_path)}"
                )
                
                doc = Document(file_path)
                
                # 处理每个表格
                for table in doc.tables:
                    # 设置表格对齐
                    if table_align == "左对齐":
                        table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    elif table_align == "居中":
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    else:
                        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                    
                    # 设置表格宽度
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph_format = paragraph.paragraph_format
                                paragraph_format.space_before = Pt(0)
                                paragraph_format.space_after = Pt(0)
                    
                    # 设置表头样式
                    if table.rows and header_bold:
                        for cell in table.rows[0].cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                
                                # 设置表头对齐
                                if header_align == "左对齐":
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                elif header_align == "居中":
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # 设置边框
                    if border_width > 0:
                        self.set_table_border(table, border_width)
                
                # 保存文档
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                new_path = os.path.join(save_dir, f"{base_name}_样式修改.docx")
                doc.save(new_path)
            
            # 完成
            if self.operation.is_running:
                message = f"表格样式修改完成，共处理 {total_files} 个文件"
                self.operation.signals.finished.emit(True, message)
            else:
                self.operation.signals.finished.emit(False, "用户取消操作")
                
        except Exception as e:
            self.operation.signals.finished.emit(False, f"修改表格样式过程中出错: {str(e)}")
            
    def process_table_operations(self, files, operation_id, position, count, save_dir):
        """处理表格操作"""
        try:
            total_files = len(files)
            
            for i, file_path in enumerate(files):
                if not self.operation.is_running:
                    break
                    
                self.operation.signals.progress.emit(
                    int(i / total_files * 90), 
                    100, 
                    f"正在处理表格: {os.path.basename(file_path)}"
                )
                
                doc = Document(file_path)
                
                # 处理每个表格
                for table in doc.tables:
                    if operation_id == 0:  # 添加列
                        self.add_columns(table, position, count)
                    elif operation_id == 1:  # 添加行
                        self.add_rows(table, position, count)
                    elif operation_id == 2:  # 删除列
                        self.delete_columns(table, position, count)
                    elif operation_id == 3:  # 删除行
                        self.delete_rows(table, position, count)
                
                # 保存文档
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                new_path = os.path.join(save_dir, f"{base_name}_表格操作.docx")
                doc.save(new_path)
            
            # 完成
            if self.operation.is_running:
                message = f"表格操作完成，共处理 {total_files} 个文件"
                self.operation.signals.finished.emit(True, message)
            else:
                self.operation.signals.finished.emit(False, "用户取消操作")
                
        except Exception as e:
            self.operation.signals.finished.emit(False, f"表格操作过程中出错: {str(e)}")
            
    def save_to_excel(self, all_data, merge_mode, save_dir, default_ext):
        """保存到Excel文件"""
        if merge_mode == "所有表格合并":
            # 所有表格合并到一个Excel文件的不同工作表
            with pd.ExcelWriter(os.path.join(save_dir, f"合并表格{default_ext}")) as writer:
                for i, table_info in enumerate(all_data):
                    df = pd.DataFrame(table_info["data"])
                    sheet_name = f"{table_info['filename']}_表格{table_info['table_index']}"[:31]  # Excel工作表名限制
                    df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)
        elif merge_mode == "每个表格单独保存":
            # 每个表格单独保存为一个Excel文件
            for table_info in all_data:
                df = pd.DataFrame(table_info["data"])
                file_name = f"{table_info['filename']}_表格{table_info['table_index']}{default_ext}"
                file_path = os.path.join(save_dir, file_name)
                df.to_excel(file_path, index=False, header=False)
        else:  # 按文档分组保存
            # 按文档分组，每个文档中的所有表格保存到一个Excel文件的不同工作表
            files_data = {}
            for table_info in all_data:
                file_name = table_info['filename']
                if file_name not in files_data:
                    files_data[file_name] = []
                files_data[file_name].append(table_info)
            
            for file_name, tables in files_data.items():
                base_name = os.path.splitext(file_name)[0]
                with pd.ExcelWriter(os.path.join(save_dir, f"{base_name}{default_ext}")) as writer:
                    for table_info in tables:
                        df = pd.DataFrame(table_info["data"])
                        sheet_name = f"表格{table_info['table_index']}"
                        df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)
                        
    def save_to_csv(self, all_data, merge_mode, save_dir, default_ext):
        """保存到CSV文件"""
        if merge_mode == "所有表格合并":
            # 所有表格合并到一个CSV文件
            merged_data = []
            for table_info in all_data:
                merged_data.extend(table_info["data"])
            
            df = pd.DataFrame(merged_data)
            df.to_csv(os.path.join(save_dir, f"合并表格{default_ext}"), index=False, header=False)
        elif merge_mode == "每个表格单独保存":
            # 每个表格单独保存为一个CSV文件
            for table_info in all_data:
                df = pd.DataFrame(table_info["data"])
                file_name = f"{table_info['filename']}_表格{table_info['table_index']}{default_ext}"
                file_path = os.path.join(save_dir, file_name)
                df.to_csv(file_path, index=False, header=False)
        else:  # 按文档分组保存
            # 按文档分组，每个文档中的所有表格保存到一个CSV文件
            files_data = {}
            for table_info in all_data:
                file_name = table_info['filename']
                if file_name not in files_data:
                    files_data[file_name] = []
                files_data[file_name].extend(table_info["data"])
            
            for file_name, data in files_data.items():
                base_name = os.path.splitext(file_name)[0]
                df = pd.DataFrame(data)
                df.to_csv(os.path.join(save_dir, f"{base_name}{default_ext}"), index=False, header=False)
                
    def save_to_json(self, all_data, merge_mode, save_dir, default_ext):
        """保存到JSON文件"""
        import json
        
        if merge_mode == "所有表格合并":
            # 所有表格合并到一个JSON文件
            with open(os.path.join(save_dir, f"合并表格{default_ext}"), 'w', encoding='utf-8') as f:
                json.dump(all_data, f, ensure_ascii=False, indent=2)
        elif merge_mode == "每个表格单独保存":
            # 每个表格单独保存为一个JSON文件
            for table_info in all_data:
                file_name = f"{table_info['filename']}_表格{table_info['table_index']}{default_ext}"
                file_path = os.path.join(save_dir, file_name)
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(table_info, f, ensure_ascii=False, indent=2)
        else:  # 按文档分组保存
            # 按文档分组，每个文档中的所有表格保存到一个JSON文件
            files_data = {}
            for table_info in all_data:
                file_name = table_info['filename']
                if file_name not in files_data:
                    files_data[file_name] = []
                files_data[file_name].append(table_info)
            
            for file_name, data in files_data.items():
                base_name = os.path.splitext(file_name)[0]
                with open(os.path.join(save_dir, f"{base_name}{default_ext}"), 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                    
    def set_table_border(self, table, width_pt):
        """设置表格边框"""
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            
            # 创建边框元素
            borders = OxmlElement('w:tblBorders')
            
            # 设置边框样式
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:sz'), str(width_pt * 2))  # sz单位是半磅
                border.set(qn('w:val'), 'single')
                border.set(qn('w:color'), '000000')
                borders.append(border)
            
            # 添加边框到表格属性
            tblPr.append(borders)
            
        except Exception as e:
            print(f"设置表格边框失败: {str(e)}")
            
    def add_columns(self, table, position, count):
        """添加列"""
        for _ in range(count):
            # 调整位置（从0开始）
            col_index = max(0, position - 1)
            
            # 为每行添加单元格
            for row in table.rows:
                # 如果位置超出了当前列数，则添加到末尾
                if col_index >= len(row.cells):
                    row.cells.add_cell()
                else:
                    # 在指定位置插入
                    from docx.oxml.table import CT_Tc, CT_Row
                    tc = CT_Tc()
                    tc_content = tc.append(OxmlElement('w:p'))
                    row._tr.insert(col_index, tc)
                    
                    # 更新行中的单元格引用
                    row.cells.insert(col_index, row.cells[col_index])
                    
    def add_rows(self, table, position, count):
        """添加行"""
        for _ in range(count):
            # 调整位置（从0开始）
            row_index = max(0, position - 1)
            
            # 创建新行
            if table.rows:
                # 复制第一行的结构
                template_row = table.rows[0]
                new_row = table.add_row()._tr
                
                # 清空单元格内容
                for cell in new_row.tc_lst:
                    for paragraph in cell.xpath('.//w:p'):
                        cell.remove(paragraph)
                    cell.append(OxmlElement('w:p'))
            else:
                # 空表格，添加空行
                table.add_row()
                
    def delete_columns(self, table, position, count):
        """删除列"""
        if not table.rows:
            return
            
        # 调整位置（从0开始）
        col_index = max(0, position - 1)
        
        for _ in range(count):
            # 检查列是否存在
            if col_index >= len(table.rows[0].cells):
                break
                
            # 删除每行的指定列
            for row in table.rows:
                if col_index < len(row.cells):
                    row._tr.remove(row.cells[col_index]._tc)
                    
    def delete_rows(self, table, position, count):
        """删除行"""
        # 调整位置（从0开始）
        row_index = max(0, position - 1)
        
        for _ in range(count):
            # 检查行是否存在
            if row_index >= len(table.rows):
                break
                
            # 删除指定行
            table._tbl.remove(table.rows[row_index]._tr)
            
    def update_progress(self, current, total, status):
        """更新进度"""
        self.operation.progress.update_progress(current, total, status)
        
    def table_processing_finished(self, success, message):
        """表格处理完成"""
        self.operation.set_running_state(False)
        if success:
            self.show_info(message)
        else:
            self.show_error(message)


if __name__ == "__main__":
    from PySide6.QtWidgets import QApplication
    app = QApplication(sys.argv)
    window = WordTableWindow()
    window.show()
    sys.exit(app.exec())