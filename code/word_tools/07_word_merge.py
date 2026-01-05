#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word文档合并工具
功能：合并多个Word文档到一个文档中，支持多种合并模式和格式保持
"""

import sys
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PySide6.QtWidgets import (QVBoxLayout, QHBoxLayout, QLabel, QComboBox, 
                              QPushButton, QCheckBox, QRadioButton, QButtonGroup,
                              QSpinBox, QMessageBox)
from PySide6.QtCore import QTimer

# 添加父目录到路径，以便导入common模块
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
sys.path.append(parent_dir)

from common.ui_base import BaseMainWindow, FileSelectionWidget, OperationWidget


class WordMergeWindow(BaseMainWindow):
    """Word文档合并工具主窗口"""
    
    def __init__(self):
        super().__init__("Word文档合并工具", 900, 650)
        
    def setup_content(self):
        """设置内容区域"""
        # 文件选择组件
        self.file_selector = FileSelectionWidget("选择要合并的Word文件", multi_selection=True)
        self.content_layout.addWidget(self.file_selector)
        
        # 合并选项组件
        self.merge_options = MergeOptionsWidget()
        self.content_layout.addWidget(self.merge_options)
        
        # 操作区域
        self.operation = OperationWidget("合并操作")
        self.operation.start_operation = self.start_merge
        self.operation.signals.progress.connect(self.update_progress)
        self.operation.signals.finished.connect(self.merge_finished)
        self.content_layout.addWidget(self.operation)
        
    def start_merge(self):
        """开始合并操作"""
        files = self.file_selector.get_files()
        if not files:
            self.show_warning("请先选择要合并的Word文件")
            return
            
        # 获取合并选项
        options = self.merge_options.get_merge_options()
        
        # 获取保存路径
        save_path = self.get_save_path("保存合并后的Word文件", "Word文件 (*.docx)", "合并文档.docx")
        if not save_path:
            return
            
        # 开始处理
        self.operation.set_running_state(True)
        self.operation.signals.progress.emit(0, 100, "开始合并Word文档...")
        
        # 使用定时器异步执行
        QTimer.singleShot(100, lambda: self.process_files(
            files, options, save_path
        ))
        
    def process_files(self, files, options, save_path):
        """处理文件合并"""
        try:
            total_files = len(files)
            
            # 创建新文档
            merged_doc = Document()
            
            # 添加标题（如果需要）
            if options["add_title"] and options["title_text"]:
                title = merged_doc.add_heading(options["title_text"], level=0)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if options["center_title"] else WD_PARAGRAPH_ALIGNMENT.LEFT
                
                # 添加标题后空行
                if options["space_after_title"] > 0:
                    for _ in range(options["space_after_title"]):
                        merged_doc.add_paragraph()
            
            # 逐个合并文档
            for i, file_path in enumerate(files):
                if not self.operation.is_running:
                    break
                    
                self.operation.signals.progress.emit(
                    int(i / total_files * 80), 
                    100, 
                    f"正在合并: {os.path.basename(file_path)}"
                )
                
                try:
                    # 打开文档
                    doc = Document(file_path)
                    
                    # 添加分隔符（除了第一个文档）
                    if i > 0 and options["add_separator"]:
                        # 添加分页符或分隔线
                        if options["separator_type"] == "page_break":
                            merged_doc.add_page_break()
                        else:
                            # 添加分隔线
                            p = merged_doc.add_paragraph()
                            p.add_run("-" * 50).bold = True
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            merged_doc.add_paragraph()
                    
                    # 添加文档标题（如果需要）
                    if options["add_doc_title"]:
                        doc_title = os.path.splitext(os.path.basename(file_path))[0]
                        doc_heading = merged_doc.add_heading(doc_title, level=1)
                        if options["doc_title_center"]:
                            doc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        # 添加文档标题后空行
                        if options["space_after_doc_title"] > 0:
                            for _ in range(options["space_after_doc_title"]):
                                merged_doc.add_paragraph()
                    
                    # 合并内容
                    if options["merge_mode"] == "append":
                        # 简单追加所有内容
                        for element in doc.element.body:
                            merged_doc.element.body.append(element)
                    elif options["merge_mode"] == "paragraph":
                        # 按段落合并
                        for paragraph in doc.paragraphs:
                            new_para = merged_doc.add_paragraph()
                            new_para._element.append(paragraph._element)
                            
                            # 应用缩进（如果需要）
                            if options["first_line_indent"] > 0:
                                new_para.paragraph_format.first_line_indent = Pt(options["first_line_indent"])
                                
                    # 合并表格
                    if options["include_tables"]:
                        for table in doc.tables:
                            merged_doc.add_paragraph()  # 表格前添加空行
                            new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                            
                            # 复制表格内容和样式
                            for i, row in enumerate(table.rows):
                                for j, cell in enumerate(row.cells):
                                    new_table.cell(i, j).text = cell.text
                                    
                    # 添加文档后空行
                    if options["space_after_doc"] > 0:
                        for _ in range(options["space_after_doc"]):
                            merged_doc.add_paragraph()
                            
                except Exception as e:
                    self.operation.signals.finished.emit(False, f"处理文件 {os.path.basename(file_path)} 时出错: {str(e)}")
                    return
            
            # 保存合并后的文档
            self.operation.signals.progress.emit(80, 100, "正在保存合并文档...")
            merged_doc.save(save_path)
            
            # 完成
            if self.operation.is_running:
                message = f"合并完成，共处理 {total_files} 个Word文档"
                self.operation.signals.finished.emit(True, message)
            else:
                self.operation.signals.finished.emit(False, "用户取消操作")
                
        except Exception as e:
            self.operation.signals.finished.emit(False, f"合并过程中发生错误: {str(e)}")
            
    def update_progress(self, current, total, status):
        """更新进度"""
        self.operation.progress.update_progress(current, total, status)
        
    def merge_finished(self, success, message):
        """合并完成"""
        self.operation.set_running_state(False)
        if success:
            self.show_info(message)
        else:
            self.show_error(message)


class MergeOptionsWidget(QWidget):
    """合并选项组件"""
    
    def __init__(self):
        super().__init__()
        self.setup_ui()
        
    def setup_ui(self):
        layout = QVBoxLayout(self)
        
        # 合并模式
        mode_layout = QHBoxLayout()
        mode_layout.addWidget(QLabel("合并模式:"))
        self.merge_mode_combo = QComboBox()
        self.merge_mode_combo.addItems(["简单追加", "段落合并"])
        mode_layout.addWidget(self.merge_mode_combo)
        layout.addLayout(mode_layout)
        
        # 主标题设置
        title_group = QWidget()
        title_layout = QVBoxLayout(title_group)
        
        self.add_title_checkbox = QCheckBox("添加主标题")
        self.add_title_checkbox.toggled.connect(self.toggle_title_options)
        title_layout.addWidget(self.add_title_checkbox)
        
        title_input_layout = QHBoxLayout()
        title_input_layout.addWidget(QLabel("标题文本:"))
        self.title_input = QLineEdit("合并文档")
        title_input_layout.addWidget(self.title_input)
        title_layout.addLayout(title_input_layout)
        
        self.center_title_checkbox = QCheckBox("标题居中")
        title_layout.addWidget(self.center_title_checkbox)
        
        space_layout = QHBoxLayout()
        space_layout.addWidget(QLabel("标题后空行数:"))
        self.space_after_title_spin = QSpinBox()
        self.space_after_title_spin.setMinimum(0)
        self.space_after_title_spin.setMaximum(10)
        self.space_after_title_spin.setValue(2)
        space_layout.addWidget(self.space_after_title_spin)
        title_layout.addLayout(space_layout)
        
        layout.addWidget(title_group)
        
        # 文档标题设置
        doc_title_group = QWidget()
        doc_title_layout = QVBoxLayout(doc_title_group)
        
        self.add_doc_title_checkbox = QCheckBox("添加每个文档的文件名作为标题")
        self.add_doc_title_checkbox.toggled.connect(self.toggle_doc_title_options)
        self.add_doc_title_checkbox.setChecked(True)
        doc_title_layout.addWidget(self.add_doc_title_checkbox)
        
        self.doc_title_center_checkbox = QCheckBox("文档标题居中")
        doc_title_layout.addWidget(self.doc_title_center_checkbox)
        
        doc_space_layout = QHBoxLayout()
        doc_space_layout.addWidget(QLabel("文档标题后空行数:"))
        self.space_after_doc_title_spin = QSpinBox()
        self.space_after_doc_title_spin.setMinimum(0)
        self.space_after_doc_title_spin.setMaximum(10)
        self.space_after_doc_title_spin.setValue(1)
        doc_space_layout.addWidget(self.space_after_doc_title_spin)
        doc_title_layout.addLayout(doc_space_layout)
        
        layout.addWidget(doc_title_group)
        
        # 分隔符设置
        separator_group = QWidget()
        separator_layout = QVBoxLayout(separator_group)
        
        self.add_separator_checkbox = QCheckBox("在文档间添加分隔符")
        self.add_separator_checkbox.setChecked(True)
        separator_layout.addWidget(self.add_separator_checkbox)
        
        separator_type_layout = QHBoxLayout()
        separator_type_layout.addWidget(QLabel("分隔符类型:"))
        self.separator_type_combo = QComboBox()
        self.separator_type_combo.addItems(["分页符", "分隔线"])
        separator_type_layout.addWidget(self.separator_type_combo)
        separator_layout.addLayout(separator_type_layout)
        
        layout.addWidget(separator_group)
        
        # 其他选项
        other_group = QWidget()
        other_layout = QVBoxLayout(other_group)
        
        self.include_tables_checkbox = QCheckBox("包含表格")
        self.include_tables_checkbox.setChecked(True)
        other_layout.addWidget(self.include_tables_checkbox)
        
        doc_space_layout = QHBoxLayout()
        doc_space_layout.addWidget(QLabel("每个文档后空行数:"))
        self.space_after_doc_spin = QSpinBox()
        self.space_after_doc_spin.setMinimum(0)
        self.space_after_doc_spin.setMaximum(10)
        self.space_after_doc_spin.setValue(1)
        doc_space_layout.addWidget(self.space_after_doc_spin)
        other_layout.addLayout(doc_space_layout)
        
        indent_layout = QHBoxLayout()
        indent_layout.addWidget(QLabel("首行缩进(磅):"))
        self.first_line_indent_spin = QSpinBox()
        self.first_line_indent_spin.setMinimum(0)
        self.first_line_indent_spin.setMaximum(72)
        self.first_line_indent_spin.setValue(24)
        indent_layout.addWidget(self.first_line_indent_spin)
        other_layout.addLayout(indent_layout)
        
        layout.addWidget(other_group)
        
        # 初始化控件状态
        self.toggle_title_options(self.add_title_checkbox.isChecked())
        self.toggle_doc_title_options(self.add_doc_title_checkbox.isChecked())
        
    def toggle_title_options(self, enabled):
        """切换标题选项可用性"""
        self.title_input.setEnabled(enabled)
        self.center_title_checkbox.setEnabled(enabled)
        self.space_after_title_spin.setEnabled(enabled)
        
    def toggle_doc_title_options(self, enabled):
        """切换文档标题选项可用性"""
        self.doc_title_center_checkbox.setEnabled(enabled)
        self.space_after_doc_title_spin.setEnabled(enabled)
        
    def get_merge_options(self):
        """获取合并选项"""
        return {
            "merge_mode": "append" if self.merge_mode_combo.currentIndex() == 0 else "paragraph",
            "add_title": self.add_title_checkbox.isChecked(),
            "title_text": self.title_input.text(),
            "center_title": self.center_title_checkbox.isChecked(),
            "space_after_title": self.space_after_title_spin.value(),
            "add_doc_title": self.add_doc_title_checkbox.isChecked(),
            "doc_title_center": self.doc_title_center_checkbox.isChecked(),
            "space_after_doc_title": self.space_after_doc_title_spin.value(),
            "add_separator": self.add_separator_checkbox.isChecked(),
            "separator_type": "page_break" if self.separator_type_combo.currentIndex() == 0 else "line",
            "include_tables": self.include_tables_checkbox.isChecked(),
            "space_after_doc": self.space_after_doc_spin.value(),
            "first_line_indent": self.first_line_indent_spin.value()
        }


if __name__ == "__main__":
    from PySide6.QtWidgets import QApplication
    app = QApplication(sys.argv)
    window = WordMergeWindow()
    window.show()
    sys.exit(app.exec())