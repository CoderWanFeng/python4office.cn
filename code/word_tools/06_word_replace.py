#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word文档批量替换工具
功能：批量替换Word文档中的文本，支持普通文本替换、格式保留和批量处理
"""

import sys
import os
from docx import Document
from docx.shared import Inches
from PySide6.QtWidgets import (QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, 
                              QPushButton, QTableWidget, QTableWidgetItem, 
                              QHeaderView, QCheckBox, QMessageBox)
from PySide6.QtCore import QTimer, Qt

# 添加父目录到路径，以便导入common模块
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
sys.path.append(parent_dir)

from common.ui_base import BaseMainWindow, FileSelectionWidget, OperationWidget


class WordReplaceWindow(BaseMainWindow):
    """Word文档批量替换工具主窗口"""
    
    def __init__(self):
        super().__init__("Word文档批量替换工具", 900, 650)
        
    def setup_content(self):
        """设置内容区域"""
        # 文件选择组件
        self.file_selector = FileSelectionWidget("选择要替换文本的Word文件", multi_selection=True)
        self.content_layout.addWidget(self.file_selector)
        
        # 替换规则组件
        self.replace_rules = ReplaceRulesWidget()
        self.content_layout.addWidget(self.replace_rules)
        
        # 替换选项组件
        self.replace_options = ReplaceOptionsWidget()
        self.content_layout.addWidget(self.replace_options)
        
        # 操作区域
        self.operation = OperationWidget("替换操作")
        self.operation.start_operation = self.start_replace
        self.operation.signals.progress.connect(self.update_progress)
        self.operation.signals.finished.connect(self.replace_finished)
        self.content_layout.addWidget(self.operation)
        
    def start_replace(self):
        """开始替换操作"""
        files = self.file_selector.get_files()
        if not files:
            self.show_warning("请先选择要处理的Word文件")
            return
            
        # 获取替换规则
        rules = self.replace_rules.get_replace_rules()
        if not rules:
            self.show_warning("请至少添加一个替换规则")
            return
            
        # 获取替换选项
        options = self.replace_options.get_replace_options()
        
        # 获取保存路径
        save_dir = self.get_directory("选择保存替换后文件的目录")
        if not save_dir:
            return
            
        # 开始处理
        self.operation.set_running_state(True)
        self.operation.signals.progress.emit(0, 100, "开始批量替换...")
        
        # 使用定时器异步执行
        QTimer.singleShot(100, lambda: self.process_files(
            files, rules, options, save_dir
        ))
        
    def process_files(self, files, rules, options, save_dir):
        """处理文件替换"""
        try:
            total_files = len(files)
            processed_files = 0
            
            for i, file_path in enumerate(files):
                if not self.operation.is_running:
                    break
                    
                self.operation.signals.progress.emit(
                    int(i / total_files * 90), 
                    100, 
                    f"正在处理文件: {os.path.basename(file_path)}"
                )
                
                try:
                    # 打开Word文档
                    doc = Document(file_path)
                    
                    # 记录是否有更改
                    has_changes = False
                    
                    # 处理段落文本
                    for paragraph in doc.paragraphs:
                        for rule in rules:
                            find_text = rule["find"]
                            replace_text = rule["replace"]
                            match_case = options["match_case"]
                            whole_word = options["whole_word"]
                            
                            if find_text in paragraph.text:
                                # 替换文本
                                if whole_word:
                                    # 全词匹配
                                    words = paragraph.text.split()
                                    new_words = []
                                    for word in words:
                                        # 简单处理标点符号
                                        clean_word = word.strip('.,!?;:"\'()[]{}')
                                        if match_case:
                                            if clean_word == find_text:
                                                new_word = word.replace(clean_word, replace_text)
                                            else:
                                                new_word = word
                                        else:
                                            if clean_word.lower() == find_text.lower():
                                                new_word = word.replace(clean_word, replace_text)
                                            else:
                                                new_word = word
                                        new_words.append(new_word)
                                    paragraph.text = ' '.join(new_words)
                                else:
                                    # 简单文本替换
                                    if match_case:
                                        if find_text in paragraph.text:
                                            paragraph.text = paragraph.text.replace(find_text, replace_text)
                                    else:
                                        # 不区分大小写替换
                                        lower_text = paragraph.text.lower()
                                        lower_find = find_text.lower()
                                        if lower_find in lower_text:
                                            # 找到所有匹配位置并替换
                                            indices = []
                                            start = 0
                                            while True:
                                                idx = lower_text.find(lower_find, start)
                                                if idx == -1:
                                                    break
                                                indices.append(idx)
                                                start = idx + len(lower_find)
                                            
                                            # 从后往前替换，避免位置偏移
                                            for idx in reversed(indices):
                                                paragraph.text = (
                                                    paragraph.text[:idx] + 
                                                    replace_text + 
                                                    paragraph.text[idx+len(find_text):]
                                                )
                                
                                has_changes = True
                    
                    # 处理表格文本
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    for rule in rules:
                                        find_text = rule["find"]
                                        replace_text = rule["replace"]
                                        match_case = options["match_case"]
                                        whole_word = options["whole_word"]
                                        
                                        if find_text in paragraph.text:
                                            if whole_word:
                                                # 全词匹配
                                                words = paragraph.text.split()
                                                new_words = []
                                                for word in words:
                                                    clean_word = word.strip('.,!?;:"\'()[]{}')
                                                    if match_case:
                                                        if clean_word == find_text:
                                                            new_word = word.replace(clean_word, replace_text)
                                                        else:
                                                            new_word = word
                                                    else:
                                                        if clean_word.lower() == find_text.lower():
                                                            new_word = word.replace(clean_word, replace_text)
                                                        else:
                                                            new_word = word
                                                    new_words.append(new_word)
                                                paragraph.text = ' '.join(new_words)
                                            else:
                                                # 简单文本替换
                                                if match_case:
                                                    if find_text in paragraph.text:
                                                        paragraph.text = paragraph.text.replace(find_text, replace_text)
                                                else:
                                                    # 不区分大小写替换
                                                    lower_text = paragraph.text.lower()
                                                    lower_find = find_text.lower()
                                                    if lower_find in lower_text:
                                                        indices = []
                                                        start = 0
                                                        while True:
                                                            idx = lower_text.find(lower_find, start)
                                                            if idx == -1:
                                                                break
                                                            indices.append(idx)
                                                            start = idx + len(lower_find)
                                                        
                                                        for idx in reversed(indices):
                                                            paragraph.text = (
                                                                paragraph.text[:idx] + 
                                                                replace_text + 
                                                                paragraph.text[idx+len(find_text):]
                                                            )
                                            
                                            has_changes = True
                    
                    # 如果有更改，保存文档
                    if has_changes or options["save_all"]:
                        # 生成新文件名
                        base_name = os.path.basename(file_path)
                        name, ext = os.path.splitext(base_name)
                        if options["add_suffix"]:
                            new_name = f"{name}_替换后{ext}"
                        else:
                            new_name = base_name
                        new_path = os.path.join(save_dir, new_name)
                        
                        # 保存文档
                        doc.save(new_path)
                        processed_files += 1
                        
                except Exception as e:
                    self.operation.signals.finished.emit(False, f"处理文件 {os.path.basename(file_path)} 时出错: {str(e)}")
                    return
            
            # 完成
            if self.operation.is_running:
                message = f"替换完成，共处理 {processed_files}/{total_files} 个文件"
                self.operation.signals.finished.emit(True, message)
            else:
                self.operation.signals.finished.emit(False, "用户取消操作")
                
        except Exception as e:
            self.operation.signals.finished.emit(False, f"替换过程中发生错误: {str(e)}")
            
    def update_progress(self, current, total, status):
        """更新进度"""
        self.operation.progress.update_progress(current, total, status)
        
    def replace_finished(self, success, message):
        """替换完成"""
        self.operation.set_running_state(False)
        if success:
            self.show_info(message)
        else:
            self.show_error(message)


class ReplaceRulesWidget(QWidget):
    """替换规则组件"""
    
    def __init__(self):
        super().__init__()
        self.setup_ui()
        
    def setup_ui(self):
        layout = QVBoxLayout(self)
        
        # 替换规则表格
        layout.addWidget(QLabel("替换规则:"))
        self.rules_table = QTableWidget(0, 2)
        self.rules_table.setHorizontalHeaderLabels(["查找文本", "替换为"])
        self.rules_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.rules_table)
        
        # 按钮区域
        btn_layout = QHBoxLayout()
        
        self.add_rule_btn = QPushButton("添加规则")
        self.add_rule_btn.clicked.connect(self.add_rule)
        btn_layout.addWidget(self.add_rule_btn)
        
        self.remove_rule_btn = QPushButton("删除规则")
        self.remove_rule_btn.clicked.connect(self.remove_rule)
        btn_layout.addWidget(self.remove_rule_btn)
        
        self.clear_rules_btn = QPushButton("清空规则")
        self.clear_rules_btn.clicked.connect(self.clear_rules)
        btn_layout.addWidget(self.clear_rules_btn)
        
        layout.addLayout(btn_layout)
        
        # 添加一个默认规则
        self.add_rule()
        
    def add_rule(self):
        """添加一条替换规则"""
        row = self.rules_table.rowCount()
        self.rules_table.insertRow(row)
        self.rules_table.setItem(row, 0, QTableWidgetItem(""))
        self.rules_table.setItem(row, 1, QTableWidgetItem(""))
        
    def remove_rule(self):
        """删除当前选中的规则"""
        current_row = self.rules_table.currentRow()
        if current_row >= 0:
            self.rules_table.removeRow(current_row)
            
    def clear_rules(self):
        """清空所有规则"""
        self.rules_table.setRowCount(0)
        
    def get_replace_rules(self):
        """获取所有替换规则"""
        rules = []
        for row in range(self.rules_table.rowCount()):
            find_item = self.rules_table.item(row, 0)
            replace_item = self.rules_table.item(row, 1)
            
            if find_item and find_item.text().strip():
                rule = {
                    "find": find_item.text().strip(),
                    "replace": replace_item.text().strip() if replace_item else ""
                }
                rules.append(rule)
                
        return rules


class ReplaceOptionsWidget(QWidget):
    """替换选项组件"""
    
    def __init__(self):
        super().__init__()
        self.setup_ui()
        
    def setup_ui(self):
        layout = QVBoxLayout(self)
        
        # 替换选项
        self.match_case_checkbox = QCheckBox("区分大小写")
        layout.addWidget(self.match_case_checkbox)
        
        self.whole_word_checkbox = QCheckBox("全词匹配")
        layout.addWidget(self.whole_word_checkbox)
        
        self.save_all_checkbox = QCheckBox("保存所有文件（包括未修改的）")
        self.save_all_checkbox.setChecked(True)
        layout.addWidget(self.save_all_checkbox)
        
        # 文件命名选项
        naming_layout = QHBoxLayout()
        self.add_suffix_checkbox = QCheckBox("添加文件名后缀")
        self.add_suffix_checkbox.setChecked(True)
        naming_layout.addWidget(self.add_suffix_checkbox)
        
        self.suffix_input = QLineEdit("_替换后")
        naming_layout.addWidget(self.suffix_input)
        
        layout.addLayout(naming_layout)
        
        # 连接信号
        self.add_suffix_checkbox.toggled.connect(self.suffix_input.setEnabled)
        
    def get_replace_options(self):
        """获取替换选项"""
        return {
            "match_case": self.match_case_checkbox.isChecked(),
            "whole_word": self.whole_word_checkbox.isChecked(),
            "save_all": self.save_all_checkbox.isChecked(),
            "add_suffix": self.add_suffix_checkbox.isChecked(),
            "suffix": self.suffix_input.text()
        }


if __name__ == "__main__":
    from PySide6.QtWidgets import QApplication
    app = QApplication(sys.argv)
    window = WordReplaceWindow()
    window.show()
    sys.exit(app.exec())