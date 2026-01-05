#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word报告生成工具
功能：基于模板和数据生成Word报告，支持多种模板类型和自定义内容
"""

import sys
import os
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from PySide6.QtWidgets import (QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, 
                              QPushButton, QTableWidget, QTableWidgetItem, 
                              QHeaderView, QTabWidget, QTextEdit, QComboBox,
                              QFileDialog, QMessageBox)
from PySide6.QtCore import QTimer

# 添加父目录到路径，以便导入common模块
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
sys.path.append(parent_dir)

from common.ui_base import BaseMainWindow, FileSelectionWidget, OperationWidget


class WordReportWindow(BaseMainWindow):
    """Word报告生成工具主窗口"""
    
    def __init__(self):
        super().__init__("Word报告生成工具", 1000, 700)
        
    def setup_content(self):
        """设置内容区域"""
        # 使用选项卡组织界面
        self.tab_widget = QTabWidget()
        self.content_layout.addWidget(self.tab_widget)
        
        # 模板选择选项卡
        self.template_tab = QWidget()
        self.setup_template_tab()
        self.tab_widget.addTab(self.template_tab, "模板与数据")
        
        # 报告内容选项卡
        self.content_tab = QWidget()
        self.setup_content_tab()
        self.tab_widget.addTab(self.content_tab, "报告内容")
        
        # 操作区域
        self.operation = OperationWidget("报告生成操作")
        self.operation.start_operation = self.generate_report
        self.operation.signals.progress.connect(self.update_progress)
        self.operation.signals.finished.connect(self.report_finished)
        self.content_layout.addWidget(self.operation)
        
    def setup_template_tab(self):
        """设置模板选项卡"""
        layout = QVBoxLayout(self.template_tab)
        
        # 模板选择
        template_layout = QHBoxLayout()
        template_layout.addWidget(QLabel("报告模板:"))
        self.template_combo = QComboBox()
        self.template_combo.addItems(["通用报告", "数据分析报告", "项目进度报告", "会议纪要", "自定义模板"])
        template_layout.addWidget(self.template_combo)
        
        self.load_template_btn = QPushButton("加载模板文件")
        self.load_template_btn.clicked.connect(self.load_custom_template)
        template_layout.addWidget(self.load_template_btn)
        
        layout.addLayout(template_layout)
        
        # 数据输入区域
        layout.addWidget(QLabel("报告数据 (JSON格式):"))
        
        self.data_table = QTableWidget(0, 2)
        self.data_table.setHorizontalHeaderLabels(["字段名", "值"])
        self.data_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.data_table)
        
        # 数据操作按钮
        data_btn_layout = QHBoxLayout()
        
        self.add_data_btn = QPushButton("添加字段")
        self.add_data_btn.clicked.connect(self.add_data_field)
        data_btn_layout.addWidget(self.add_data_btn)
        
        self.remove_data_btn = QPushButton("删除字段")
        self.remove_data_btn.clicked.connect(self.remove_data_field)
        data_btn_layout.addWidget(self.remove_data_btn)
        
        self.clear_data_btn = QPushButton("清空数据")
        self.clear_data_btn.clicked.connect(self.clear_data_fields)
        data_btn_layout.addWidget(self.clear_data_btn)
        
        layout.addLayout(data_btn_layout)
        
        # 示例数据按钮
        self.load_example_btn = QPushButton("加载示例数据")
        self.load_example_btn.clicked.connect(self.load_example_data)
        layout.addWidget(self.load_example_btn)
        
        # 添加一些默认数据
        self.load_example_data()
        
    def setup_content_tab(self):
        """设置内容选项卡"""
        layout = QVBoxLayout(self.content_tab)
        
        # 报告基本信息
        info_group = QWidget()
        info_layout = QVBoxLayout(info_group)
        info_layout.addWidget(QLabel("报告基本信息:"))
        
        # 标题
        title_layout = QHBoxLayout()
        title_layout.addWidget(QLabel("报告标题:"))
        self.title_input = QLineEdit("项目分析报告")
        title_layout.addWidget(self.title_input)
        info_layout.addLayout(title_layout)
        
        # 作者
        author_layout = QHBoxLayout()
        author_layout.addWidget(QLabel("报告作者:"))
        self.author_input = QLineEdit("数据分析部")
        author_layout.addWidget(self.author_input)
        info_layout.addLayout(author_layout)
        
        # 日期
        date_layout = QHBoxLayout()
        date_layout.addWidget(QLabel("报告日期:"))
        from PySide6.QtCore import QDate
        from PySide6.QtWidgets import QDateEdit
        self.date_input = QDateEdit(QDate.currentDate())
        date_layout.addWidget(self.date_input)
        info_layout.addLayout(date_layout)
        
        layout.addWidget(info_group)
        
        # 报告内容
        layout.addWidget(QLabel("报告内容:"))
        self.content_editor = QTextEdit()
        self.content_editor.setPlaceholderText("输入报告正文内容，可以使用 {字段名} 引用数据中的字段")
        self.content_editor.setText(
            "# {项目名称}\n\n"
            "## 概述\n"
            "本报告分析了 {项目名称} 的执行情况，截至 {日期}，项目整体进展如下。\n\n"
            "## 数据分析\n"
            "### 完成进度\n"
            "项目已完成 {完成进度}%，计划完成时间为 {完成时间}。\n\n"
            "### 资源投入\n"
            "项目已投入 {资源投入} 人力资源，预算使用率达到 {预算使用率}%。\n\n"
            "## 结论\n"
            "综上所述，{项目名称} 当前状态为 {项目状态}，建议 {建议措施}。\n"
        )
        layout.addWidget(self.content_editor)
        
    def load_custom_template(self):
        """加载自定义模板文件"""
        template_path = QFileDialog.getOpenFileName(self, "选择Word模板文件", "", "Word文件 (*.docx)")[0]
        if template_path:
            self.custom_template_path = template_path
            self.template_combo.setCurrentText("自定义模板")
            self.show_info(f"已加载自定义模板: {os.path.basename(template_path)}")
            
    def add_data_field(self):
        """添加数据字段"""
        row = self.data_table.rowCount()
        self.data_table.insertRow(row)
        self.data_table.setItem(row, 0, QTableWidgetItem(""))
        self.data_table.setItem(row, 1, QTableWidgetItem(""))
        
    def remove_data_field(self):
        """删除数据字段"""
        current_row = self.data_table.currentRow()
        if current_row >= 0:
            self.data_table.removeRow(current_row)
            
    def clear_data_fields(self):
        """清空数据字段"""
        self.data_table.setRowCount(0)
        
    def load_example_data(self):
        """加载示例数据"""
        example_data = {
            "项目名称": "办公自动化系统开发",
            "日期": "2023-11-13",
            "完成进度": "75",
            "完成时间": "2023-12-31",
            "资源投入": "10人/月",
            "预算使用率": "65",
            "项目状态": "按计划进行",
            "建议措施": "继续按计划推进，关注风险点"
        }
        
        self.data_table.setRowCount(0)
        
        for key, value in example_data.items():
            row = self.data_table.rowCount()
            self.data_table.insertRow(row)
            self.data_table.setItem(row, 0, QTableWidgetItem(key))
            self.data_table.setItem(row, 1, QTableWidgetItem(str(value)))
            
    def generate_report(self):
        """生成报告"""
        # 获取数据
        data = self.get_report_data()
        if not data:
            self.show_warning("请至少添加一个数据字段")
            return
            
        # 获取报告内容
        content = self.content_editor.toPlainText()
        if not content:
            self.show_warning("请输入报告内容")
            return
            
        # 获取基本信息
        title = self.title_input.text()
        author = self.author_input.text()
        date = self.date_input.date().toString("yyyy-MM-dd")
        
        # 获取模板选择
        template_type = self.template_combo.currentText()
        
        # 获取保存路径
        save_path = self.get_save_path("保存生成的报告", "Word文件 (*.docx)", f"{title}.docx")
        if not save_path:
            return
            
        # 开始生成
        self.operation.set_running_state(True)
        self.operation.signals.progress.emit(0, 100, "开始生成报告...")
        
        # 使用定时器异步执行
        QTimer.singleShot(100, lambda: self.process_report(
            data, content, title, author, date, template_type, save_path
        ))
        
    def get_report_data(self):
        """获取报告数据"""
        data = {}
        for row in range(self.data_table.rowCount()):
            key_item = self.data_table.item(row, 0)
            value_item = self.data_table.item(row, 1)
            
            if key_item and key_item.text().strip():
                key = key_item.text().strip()
                value = value_item.text().strip() if value_item else ""
                data[key] = value
                
        return data
        
    def process_report(self, data, content, title, author, date, template_type, save_path):
        """处理报告生成"""
        try:
            # 创建文档
            if template_type == "自定义模板" and hasattr(self, 'custom_template_path'):
                # 使用自定义模板
                doc = Document(self.custom_template_path)
                # 清除模板内容（保留样式）
                for paragraph in doc.paragraphs[:]:
                    p = paragraph._element
                    p.getparent().remove(p)
            else:
                # 创建新文档
                doc = Document()
            
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # 添加标题
            self.operation.signals.progress.emit(20, 100, "正在添加报告标题...")
            title_heading = doc.add_heading(title, level=0)
            title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加报告信息
            info_paragraph = doc.add_paragraph()
            info_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            info_paragraph.add_run(f"作者: {author}    日期: {date}")
            info_paragraph.runs[0].italic = True
            
            # 添加分页符（可选）
            # doc.add_page_break()
            
            # 处理内容
            self.operation.signals.progress.emit(50, 100, "正在处理报告内容...")
            
            # 替换占位符
            processed_content = content
            for key, value in data.items():
                processed_content = processed_content.replace(f"{{{key}}}", str(value))
            
            # 解析内容
            lines = processed_content.split('\n')
            current_paragraph = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    # 空行，添加空段落
                    doc.add_paragraph()
                    continue
                
                # 检查是否是标题
                if line.startswith('# '):
                    # 一级标题
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    # 二级标题
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    # 三级标题
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    # 四级标题
                    doc.add_heading(line[5:], level=4)
                else:
                    # 普通段落
                    paragraph = doc.add_paragraph(line)
            
            # 添加页脚（可选）
            self.operation.signals.progress.emit(80, 100, "正在添加页脚...")
            self.add_footer(doc, f"{title} - {date}")
            
            # 保存文档
            self.operation.signals.progress.emit(90, 100, "正在保存报告...")
            doc.save(save_path)
            
            # 完成
            if self.operation.is_running:
                self.operation.signals.finished.emit(True, f"报告已生成: {save_path}")
            else:
                self.operation.signals.finished.emit(False, "用户取消操作")
                
        except Exception as e:
            self.operation.signals.finished.emit(False, f"生成报告过程中出错: {str(e)}")
            
    def add_footer(self, doc, text):
        """添加页脚"""
        try:
            # 获取节
            section = doc.sections[0]
            footer = section.footer
            
            # 清除现有页脚
            for paragraph in footer.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)
            
            # 添加页脚段落
            footer_para = footer.paragraphs[0]
            footer_para.text = text
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加页码
            add_page_number(footer)
            
        except Exception as e:
            print(f"添加页脚失败: {str(e)}")
            
    def update_progress(self, current, total, status):
        """更新进度"""
        self.operation.progress.update_progress(current, total, status)
        
    def report_finished(self, success, message):
        """报告生成完成"""
        self.operation.set_running_state(False)
        if success:
            self.show_info(message)
        else:
            self.show_error(message)


def add_page_number(footer):
    """添加页码到页脚"""
    try:
        from docx.oxml.parser import parse_xml
        from docx.oxml.ns import nsdecls
        
        # 创建页码段落
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加页码字段
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run = paragraph.runs[0]._r
        run.append(fldChar1)
        run.append(instrText)
        run.append(fldChar2)
        
    except Exception as e:
        print(f"添加页码失败: {str(e)}")


if __name__ == "__main__":
    from PySide6.QtWidgets import QApplication, QDate
    app = QApplication(sys.argv)
    window = WordReportWindow()
    window.show()
    sys.exit(app.exec())